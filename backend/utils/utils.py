from datetime import datetime
from io import BytesIO
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from PIL import Image
import io
import base64
import json
from utils.Extract_Data_Single_File import data_Extract, single_file
import pickle
import os
from docx import Document
from docx.shared import Inches
import matplotlib
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont

matplotlib.use('agg')

def f(x):
    if type(x) == str:
        return x.lower()


def save_docx(data):
    
    document = Document()
    # Compute available content width in inches (page width minus margins)
    try:
        section = document.sections[0]
        content_width_inches = (section.page_width - section.left_margin - section.right_margin) / 914400
    except Exception:
        # Fallback to a sensible default if section metrics are unavailable
        content_width_inches = 6.5
    
    #document.add_paragraph(data['title'])
    
    #################################
    base64_image = data['image6'][22:]
    # Decode the base64 image and load it using PIL
    image_data = BytesIO(base64.b64decode(base64_image))
    img = Image.open(image_data)
    # Calculate the width and height of the image to maintain aspect ratio
    width, height = img.size
    aspect_ratio = width / height
    # Set the desired width of the image to the full content width
    desired_width = content_width_inches  # in inches
    # Calculate the corresponding height based on the aspect ratio
    desired_height = desired_width / aspect_ratio
    # Add the image to the document with the specified width and height
    document.add_picture(image_data, width=Inches(desired_width), height=Inches(desired_height))
    ##################################

    document.add_paragraph(data["para1"])

    ################document.add_paragraph(data["paragraph1"])
    ################document.add_paragraph(data["paragraph2"])

    #################################
    base64_image = data['image4'][22:]
    # Decode the base64 image and load it using PIL
    image_data = BytesIO(base64.b64decode(base64_image))
    img = Image.open(image_data)
    # Calculate the width and height of the image to maintain aspect ratio
    width, height = img.size
    aspect_ratio = width / height
    # Set the desired width of the image to the full content width
    desired_width = content_width_inches  # in inches
    # Calculate the corresponding height based on the aspect ratio
    desired_height = desired_width / aspect_ratio
    # Add the image to the document with the specified width and height
    document.add_picture(image_data, width=Inches(desired_width), height=Inches(desired_height))
    ##################################

    #################################
    base64_image = data['image5'][22:]
    # Decode the base64 image and load it using PIL
    image_data = BytesIO(base64.b64decode(base64_image))
    img = Image.open(image_data)
    # Calculate the width and height of the image to maintain aspect ratio
    width, height = img.size
    aspect_ratio = width / height
    # Set the desired width of the image to the full content width
    desired_width = content_width_inches  # in inches
    # Calculate the corresponding height based on the aspect ratio
    desired_height = desired_width / aspect_ratio
    # Add the image to the document with the specified width and height
    document.add_picture(image_data, width=Inches(desired_width), height=Inches(desired_height))
    ##################################


    # document.add_paragraph()
    # document.add_paragraph()
    # document.add_paragraph()
    # document.add_paragraph()
    # document.add_paragraph()

    #################################
    base64_image = data['image8'][22:]
    # Decode the base64 image and load it using PIL
    image_data = BytesIO(base64.b64decode(base64_image))
    img = Image.open(image_data)
    # Calculate the width and height of the image to maintain aspect ratio
    width, height = img.size
    aspect_ratio = width / height
    # Set the desired width of the image to the full content width
    desired_width = content_width_inches  # in inches
    # Calculate the corresponding height based on the aspect ratio
    desired_height = desired_width / aspect_ratio
    # Add the image to the document with the specified width and height
    document.add_picture(image_data, width=Inches(desired_width), height=Inches(desired_height))
    ##################################

    base64_image = data['image1'][22:]
    # Decode the base64 image and load it using PIL
    image_data = BytesIO(base64.b64decode(base64_image))
    img = Image.open(image_data)
    # Calculate the width and height of the image to maintain aspect ratio
    width, height = img.size
    aspect_ratio = width / height
    # Set the desired width of the image to the full content width
    desired_width = content_width_inches  # in inches
    # Calculate the corresponding height based on the aspect ratio
    desired_height = desired_width / aspect_ratio
    # Add the image to the document with the specified width and height
    document.add_picture(image_data, width=Inches(desired_width), height=Inches(desired_height))

    #################################
    base64_image = data['image3'][22:]
    # Decode the base64 image and load it using PIL
    image_data = BytesIO(base64.b64decode(base64_image))
    img = Image.open(image_data)
    # Calculate the width and height of the image to maintain aspect ratio
    width, height = img.size
    aspect_ratio = width / height
    # Set the desired width of the image to the full content width
    desired_width = content_width_inches  # in inches
    # Calculate the corresponding height based on the aspect ratio
    desired_height = desired_width / aspect_ratio
    # Add the image to the document with the specified width and height
    document.add_picture(image_data, width=Inches(desired_width), height=Inches(desired_height))
    ##################################

    ####################document.add_paragraph(data["paragraph3"])
    base64_image = data['image2'][22:]
    # Decode the base64 image and load it using PIL
    image_data = BytesIO(base64.b64decode(base64_image))
    img = Image.open(image_data)
    # Calculate the width and height of the image to maintain aspect ratio
    width, height = img.size
    aspect_ratio = width / height
    # Set the desired width of the image in the document (you can adjust this as needed)
    desired_width = 6  # in inches
    # Calculate the corresponding height based on the aspect ratio
    desired_height = desired_width / aspect_ratio
    # Add the image to the document with the specified width and height
    document.add_picture(image_data, width=Inches(desired_width), height=Inches(desired_height))
    ############################document.add_paragraph(data["paragraph4"])

    temp_filename = 'temp.docx'
    document.save(temp_filename)

def get_data(data, data_file_filename, fields_tuple_1, fields_tuple_2, handicap_tuple, toggle_button='other', arena_toggled=False):

    output = generate_image_data(data, data_file_filename, fields_tuple_1, fields_tuple_2, handicap_tuple, arena_toggled)

    base64_image_title = output[0]
    base64_image_detail = output[1] 
    base64_image_table = output[2] 
    base64_image_line = output[3] 
    base64_image_chart = output[4] 
    base64_image_matchinfo = output[5] 
    base64_image_combinedimage = output[6]
    base64_image_fields_layout = output[7]

    para1 = get_text_first_paragraph(data, data_file_filename, toggle_button)

    # Create a dictionary to hold the data
    data = {
        "image1": "data:image/png;base64,"+base64_image_table,
        "image2": "data:image/png;base64,"+base64_image_chart,
        "image3": "data:image/png;base64,"+base64_image_line,
        "image4": "data:image/png;base64,"+base64_image_title,
        "image5": "data:image/png;base64,"+base64_image_detail,
        'para1': para1,
        "image6": "data:image/png;base64,"+base64_image_matchinfo,
        "image7": "data:image/png;base64,"+base64_image_combinedimage,
        "image8": "data:image/png;base64,"+base64_image_fields_layout
    }

    return data

def generate_extra_data(data, data_file_filename):
    print(f"DEBUG: generate_extra_data called with filename: {data_file_filename}")
    print(f"DEBUG: Data shape: {data.shape}")
    print(f"DEBUG: Data columns: {list(data.columns)}")
    
    newdata = data[data["Team"] != "Team"]
    print(f"DEBUG: After filtering, newdata shape: {newdata.shape}")
    
    teams = newdata["Team"].unique()
    print(f"DEBUG: Teams found: {teams}")
    title = f"{teams[0]} Vs {teams[1]}"

    print("DEBUG: Calling data_Extract...")
    finalLst = data_Extract(data)
    print(f"DEBUG: data_Extract returned list of length: {len(finalLst)}")
    
    print("DEBUG: Calling single_file...")
    extracted_data, filename = single_file(finalLst, data_file_filename)
    print(f"DEBUG: single_file returned filename: {filename}")
    print(f"DEBUG: extracted_data columns: {list(extracted_data.columns)}")

    team1_name   = extracted_data['team1'][0]
    team2_name   = extracted_data['team2'][0]
    team1_goals  = extracted_data['goals1'][0]
    team2_goals  = extracted_data['goals2'][0]
    club         = extracted_data['club'][0]
    year         = extracted_data['year'][0]
    tournament   = extracted_data['tournament'][0]

    print(f"DEBUG: Extracted data - Team1: {team1_name}, Team2: {team2_name}")
    print(f"DEBUG: Goals - Team1: {team1_goals}, Team2: {team2_goals}")
    print(f"DEBUG: Club: {club}, Year: {year}, Tournament: {tournament}")

    # Parse date flexibly
    input_date_str = extracted_data['date'][0].strip()
    print(f"DEBUG: Date string to parse: '{input_date_str}'")
    
    for fmt in ("%B %d. %Y", "%B %d %Y", "%B %d.%Y"):
        try:
            input_date = datetime.strptime(input_date_str, fmt)
            print(f"DEBUG: Successfully parsed date with format: {fmt}")
            break
        except ValueError as e:
            print(f"DEBUG: Failed to parse with format {fmt}: {e}")
            continue
    else:
        print(f"DEBUG: All date parsing attempts failed")
        raise ValueError(
            f"Date '{input_date_str}' is not in a supported format "
            "(expected like 'January 17. 2021' or 'January 17 2021')"
        )

    date_str = input_date.strftime("%Y-%m-%d")
    print(f"DEBUG: Final date string: {date_str}")

    return (
        team1_name,
        team2_name,
        team1_goals,
        team2_goals,
        club,
        year,
        tournament,
        date_str
    )


def generate_table_data(data):
    data = data[data["Team"]!="Team"]
    teams = data["Team"].unique()
    team1 = teams[0]
    team2 = teams[1]
    data1 = data[data["Team"]==team1]
    data2 = data[data["Team"]==team2]

    data1_goal = data1.loc[data1["Goal"].dropna().index]
    goals1Chukka1 = len(data1_goal[data1_goal["Chukka"]=="Chukka 1"])
    goals1Chukka2 = len(data1_goal[data1_goal["Chukka"]=="Chukka 2"])
    goals1Chukka3 = len(data1_goal[data1_goal["Chukka"]=="Chukka 3"])
    goals1Chukka4 = len(data1_goal[data1_goal["Chukka"]=="Chukka 4"])
    goals1Chukka5 = len(data1_goal[data1_goal["Chukka"]=="Chukka 5"])
    goals1Chukka6 = len(data1_goal[data1_goal["Chukka"]=="Chukka 6"])

    data2_goal = data2.loc[data2["Goal"].dropna().index]
    goals2Chukka1 = len(data2_goal[data2_goal["Chukka"]=="Chukka 1"])
    goals2Chukka2 = len(data2_goal[data2_goal["Chukka"]=="Chukka 2"])
    goals2Chukka3 = len(data2_goal[data2_goal["Chukka"]=="Chukka 3"])
    goals2Chukka4 = len(data2_goal[data2_goal["Chukka"]=="Chukka 4"])
    goals2Chukka5 = len(data2_goal[data2_goal["Chukka"]=="Chukka 5"])
    goals2Chukka6 = len(data2_goal[data2_goal["Chukka"]=="Chukka 6"])

    json_string = {
        'Team': [team1, team2],
        'Chukka1': [goals1Chukka1, goals2Chukka1],
        'Chukka2': [goals1Chukka2, goals2Chukka2],
        'Chukka3': [goals1Chukka3, goals2Chukka3],
        'Chukka4': [goals1Chukka4, goals2Chukka4],
        'Chukka5': [goals1Chukka5, goals2Chukka5],
        'Chukka6': [goals1Chukka6, goals2Chukka6]
    }

    json_string = json.dumps(json_string)

    return json_string

def generate_bar_data(data):
    
    data = data[data["Team"]!="Team"]
    teams = data["Team"].unique()
    team1 = teams[0]
    team2 = teams[1]
    data1 = data[data["Team"]==team1]
    data2 = data[data["Team"]==team2]

    goals1 = len(data1["Goal"].dropna())
    goals2 = len(data2["Goal"].dropna())

    fouls1 = len(data1[data1["Foul"]=="Foul"])
    fouls2 = len(data2[data2["Foul"]=="Foul"])

    shotAtGoal1 = len(data1["Shot At Goal"].dropna())
    shotAtGoal2 = len(data2["Shot At Goal"].dropna())

    # shotAtGoalOutcomeGoal1 = len(data1[data1["Shot At Goal Outcome"]=="Goal"])
    # shotAtGoalOutcomeGoal2 = len(data2[data2["Shot At Goal Outcome"]=="Goal"])    
    shotAtGoalOutcomeGoal1 = len(data1[data1["Shot At Goal Outcome"]=="Goal"])
    shotAtGoalOutcomeGoal1 += len(data1[data1["Shot At Goal Outcome"]=="2pt Goal"])*2
    shotAtGoalOutcomeGoal2 = len(data2[data2["Shot At Goal Outcome"]=="Goal"])
    shotAtGoalOutcomeGoal2 += len(data2[data2["Shot At Goal Outcome"]=="2pt Goal"])*2

    penaltyShots1 = len(data1["Penalty Outcome"].dropna())
    penaltyShots2 = len(data2["Penalty Outcome"].dropna())

    penaltyShotsGoal1 = len(data1[data1["Penalty Outcome"]=="Goal"])
    penaltyShotsGoal2 = len(data2[data2["Penalty Outcome"]=="Goal"])

    throw_ins = data['Throw-in Outcome'].dropna()
    throw_ins_1 = throw_ins.str.contains(team1).sum()
    throw_ins_2 = throw_ins.str.contains(team2).sum()
    
    # Data
    categories = ['Goals', 'Field Shots', 'Field Goals', 'Penalty Shots', 'Penalty Goals', 'Throw-ins Won', 'Fouls Committed']
    col1_values = [goals1, shotAtGoal1, shotAtGoalOutcomeGoal1, penaltyShots1, penaltyShotsGoal1, throw_ins_1, fouls1]
    col2_values = [goals2, shotAtGoal2, shotAtGoalOutcomeGoal2, penaltyShots2, penaltyShotsGoal2, throw_ins_2, fouls2]

    json_string = {
        'Goals': [int(goals1), int(goals2)],
        'Field Shots': [int(shotAtGoal1), int(shotAtGoal2)],
        'Field Goals': [int(shotAtGoalOutcomeGoal1), int(shotAtGoalOutcomeGoal2)],
        'Penalty Shots': [int(penaltyShots1), int(penaltyShots2)],
        'Penalty Goals': [int(penaltyShotsGoal1), int(penaltyShotsGoal2)],
        'Throw-ins Won': [int(throw_ins_1), int(throw_ins_2)],
        'Fouls Committed': [int(fouls1), int(fouls2)]
    }

    json_string = json.dumps(json_string)

    return json_string

def generate_image_data(data, data_file_filename, fields_tuple_1, fields_tuple_2, handicap_tuple, arena_toggled):
    #data = pd.read_csv('Sample Data/Aliano Realty vs Palm Beach Equine- Live Tags.csv')

    save_table(data, handicap_tuple)
    save_line(data, handicap_tuple)
    save_barchart(data, handicap_tuple)
    save_team_goals(data, data_file_filename, handicap_tuple)
    save_game_info(data, data_file_filename)
    save_title_date(data, data_file_filename)
    #player_orders(fields_tuple_1, fields_tuple_2)
    player_orders_with_arena(fields_tuple_1, fields_tuple_2, arena_toggled)
    create_combine_image()

    image_path = "formatted_result.png"
    loaded_image = load_image(image_path)
    base64_image_title = image_to_base64(loaded_image)

    image_path = "formatted_result_game.png"
    loaded_image = load_image(image_path)
    base64_image_detail = image_to_base64(loaded_image)

    image_path = "table.png"
    loaded_image = load_image(image_path)
    base64_image_table = image_to_base64(loaded_image)

    image_path = "stacked_line_chart.png"
    loaded_image = load_image(image_path)
    base64_image_line = image_to_base64(loaded_image)    

    image_path = "stacked_bar_chart.png"
    loaded_image = load_image(image_path)
    base64_image_chart = image_to_base64(loaded_image)    

    image_path = "match_info.png"
    loaded_image = load_image(image_path)
    base64_image_matchinfo = image_to_base64(loaded_image)    

    image_path = "combined_image.png"
    loaded_image = load_image(image_path)
    base64_image_combinedimage = image_to_base64(loaded_image)    

    image_path = "fields_layout.png"
    loaded_image = load_image(image_path)
    base64_image_fields_layout = image_to_base64(loaded_image)    


    output = (base64_image_title, base64_image_detail, 
    base64_image_table, base64_image_line, 
    base64_image_chart, base64_image_matchinfo,
    base64_image_combinedimage, base64_image_fields_layout)
    
    return output

def load_image(image_path):
    image = Image.open(image_path)
    return image

def image_to_base64(image):
    image_byte_array = io.BytesIO()
    image.save(image_byte_array, format=image.format)
    encoded_string = base64.b64encode(image_byte_array.getvalue()).decode("utf-8")
    return encoded_string
        
def save_table(data, handicap_tuple):

    data = data[data["Team"]!="Team"]
    teams = data["Team"].unique()
    team1 = teams[0]
    team2 = teams[1]
    data1 = data[data["Team"]==team1]
    data2 = data[data["Team"]==team2]

    total_chukka = len(data['Chukka'].unique())

    #### Handicap data ####
    handicap_teams = [team1,team2]
    handicap_goals = [0,0]
    data['Goal_New'] = data['Goal'].map(f)
    filter = data['Goal_New'].str.contains('handicap') == True
    handicap_goal_sum = sum(filter)
    if handicap_goal_sum>0:
        handicap_goals_team = data[filter]['Team'].iloc[0]
        if team1 == handicap_goals_team:
            handicap_goals[0] = handicap_goal_sum
        else:
            handicap_goals[1] = handicap_goal_sum    
    #### Handicap data end####
    
    data1_goal = data1.loc[data1["Goal"].dropna().index]
    goals1Chukka1 = len(data1_goal[data1_goal["Chukka"]=="Chukka 1"])
    goals1Chukka2 = len(data1_goal[data1_goal["Chukka"]=="Chukka 2"])
    goals1Chukka3 = len(data1_goal[data1_goal["Chukka"]=="Chukka 3"])
    goals1Chukka4 = len(data1_goal[data1_goal["Chukka"]=="Chukka 4"])
    goals1Chukka5 = len(data1_goal[data1_goal["Chukka"]=="Chukka 5"])
    goals1Chukka6 = len(data1_goal[data1_goal["Chukka"]=="Chukka 6"])

    data2_goal = data2.loc[data2["Goal"].dropna().index]
    goals2Chukka1 = len(data2_goal[data2_goal["Chukka"]=="Chukka 1"])
    goals2Chukka2 = len(data2_goal[data2_goal["Chukka"]=="Chukka 2"])
    goals2Chukka3 = len(data2_goal[data2_goal["Chukka"]=="Chukka 3"])
    goals2Chukka4 = len(data2_goal[data2_goal["Chukka"]=="Chukka 4"])
    goals2Chukka5 = len(data2_goal[data2_goal["Chukka"]=="Chukka 5"])
    goals2Chukka6 = len(data2_goal[data2_goal["Chukka"]=="Chukka 6"])

    total_goals_1 = goals1Chukka1 + goals1Chukka2 + goals1Chukka3 + goals1Chukka4 + goals1Chukka5 + goals1Chukka6
    total_goals_2 = goals2Chukka1 + goals2Chukka2 + goals2Chukka3 + goals2Chukka4 + goals2Chukka5 + goals2Chukka6

    hg_winning, hg_losing = handicap_tuple
    hg_winning, hg_losing = float(hg_winning), float(hg_losing)
    if total_goals_2>total_goals_1:
        goals2Chukka1 = goals2Chukka1 + hg_winning
        goals1Chukka1 = goals1Chukka1 + hg_losing
    else:
        goals1Chukka1 = goals1Chukka1 + hg_winning
        goals2Chukka1 = goals2Chukka1 + hg_losing
    
    # Convert to integers if the results are whole numbers
    if isinstance(goals1Chukka1, float) and goals1Chukka1.is_integer():
        goals1Chukka1 = int(goals1Chukka1)
    if isinstance(goals2Chukka1, float) and goals2Chukka1.is_integer():
        goals2Chukka1 = int(goals2Chukka1)

    if (total_chukka == 4):
        # Create a DataFrame with column names and two rows
        cols = ['Team', 'Chukka'] + [''] * 3
        data = pd.DataFrame(columns=cols)
        data.loc[0] = [''] + [1,2,3,4]
        data.loc[1] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4]
        data.loc[2] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4]

    elif (total_chukka == 7):
        goals1Chukka7 = len(data1_goal[data1_goal["Chukka"]=="Chukka 7"])
        goals2Chukka7 = len(data2_goal[data2_goal["Chukka"]=="Chukka 7"])
    
        # Create a DataFrame with column names and two rows
        cols = ['Team', 'Chukka'] + [''] * 6
        data = pd.DataFrame(columns=cols)
        data.loc[0] = [''] + [1,2,3,4,5,6,7]
        data.loc[1] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4, goals1Chukka5, goals1Chukka6, goals1Chukka7]
        data.loc[2] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4, goals2Chukka5, goals2Chukka6, goals2Chukka7]

    elif (total_chukka == 8):
        goals1Chukka7 = len(data1_goal[data1_goal["Chukka"]=="Chukka 7"])
        goals2Chukka7 = len(data2_goal[data2_goal["Chukka"]=="Chukka 7"])
        goals1Chukka8 = len(data1_goal[data1_goal["Chukka"]=="Chukka 8"])
        goals2Chukka8 = len(data2_goal[data2_goal["Chukka"]=="Chukka 8"])
    
        # Create a DataFrame with column names and two rows
        cols = ['Team', 'Chukka'] + [''] * 7
        data = pd.DataFrame(columns=cols)
        data.loc[0] = [''] + [1,2,3,4,5,6,7,8]
        data.loc[1] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4, goals1Chukka5, goals1Chukka6, goals1Chukka7, goals1Chukka8]
        data.loc[2] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4, goals2Chukka5, goals2Chukka6, goals2Chukka7, goals2Chukka8]

    elif (total_chukka == 9):
        goals1Chukka7 = len(data1_goal[data1_goal["Chukka"]=="Chukka 7"])
        goals2Chukka7 = len(data2_goal[data2_goal["Chukka"]=="Chukka 7"])
        goals1Chukka8 = len(data1_goal[data1_goal["Chukka"]=="Chukka 8"])
        goals2Chukka8 = len(data2_goal[data2_goal["Chukka"]=="Chukka 8"])
        goals1Chukka9 = len(data1_goal[data1_goal["Chukka"]=="Chukka 9"])
        goals2Chukka9 = len(data2_goal[data2_goal["Chukka"]=="Chukka 9"])
    
        # Create a DataFrame with column names and two rows
        cols = ['Team', 'Chukka'] + [''] * 8
        data = pd.DataFrame(columns=cols)
        data.loc[0] = [''] + [1,2,3,4,5,6,7,8,9]
        data.loc[1] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4, goals1Chukka5, goals1Chukka6, goals1Chukka7, goals1Chukka8, goals1Chukka9]
        data.loc[2] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4, goals2Chukka5, goals2Chukka6, goals2Chukka7, goals2Chukka8, goals2Chukka9]

    else:        
        # Create a DataFrame with column names and two rows
        cols = ['Team', 'Chukka'] + [''] * 5
        data = pd.DataFrame(columns=cols)
        data.loc[0] = [''] + [1,2,3,4,5,6]
        data.loc[1] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4, goals1Chukka5, goals1Chukka6]
        data.loc[2] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4, goals2Chukka5, goals2Chukka6]

    # Create a table plot
    fig, ax = plt.subplots(figsize=(12, 2))

    # Remove axis
    ax.axis('off')

    # Define custom color codes
    maroon_red = '#800000'
    light_grey = '#D3D3D3'
    white = '#FFFFFF'

    data_values = data.values

    new_values = []

    # Iterate over each row in data
    for sublist in data_values:
        new_row = []

        # Iterate over each value in the sublist
        for value in sublist:
            try:
                float_value = float(value)
                if float_value.is_integer():
                    new_row.append(int(float_value))
                else:
                    new_row.append(float_value)  # If it's a float but not an integer, keep it as float
            except ValueError:
                new_row.append(value)  # If it can't be converted to float, keep the original value

        new_values.append(new_row)
    
    frac_for_chukka = 0.7
    team_name_chars = max(len(team1), len(team2))
    if team_name_chars > 30 and team_name_chars <= 40:
        frac_for_chukka = 0.55
    elif team_name_chars > 40:
        frac_for_chukka = 0.5

    if total_chukka == 4:
        frac_for_chukka = 0.5

    chukka_cols = len(data.columns) - 1
    frac = frac_for_chukka / chukka_cols
    # Allocate a bit more width to the first chukka column to ensure the
    # "Chukka" header is not truncated, then distribute the rest evenly.
    if chukka_cols > 1:
        first_chukka_width = min(frac * 1.6, frac_for_chukka * 0.6)
        remaining = max(frac_for_chukka - first_chukka_width, 0)
        other_chukka_width = remaining / (chukka_cols - 1)
        col_width = [1 - frac_for_chukka] + [first_chukka_width] + [other_chukka_width] * (chukka_cols - 1)
    else:
        # Only one chukka column
        col_width = [1 - frac_for_chukka] + [frac_for_chukka]

    # Create the table
    table = plt.table(cellText=new_values, #data.values,
                    colLabels=data.columns,
                    cellLoc='center',
                    loc='center',
                    cellColours=[[maroon_red]*len(data.columns), ['lightgrey']*len(data.columns), ['white']*len(data.columns)],
                    colColours=[maroon_red]*len(data.columns),
                    colWidths= col_width,
                    bbox=[0, 0, 1, 1])

    # Set font size and color for column headers
    table.auto_set_font_size(False)
    table.set_fontsize(14)

    for i, key in enumerate(table.get_celld().keys()):
        if key[0] == 0 or key[0] == 1:
            table[key].set_text_props(weight='bold', color='white')

    # Hide cell borders
    for key, cell in table.get_celld().items():
        cell.set_linewidth(0)

    # Hide row and column separators
    ax.spines['top'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)

    # Associate table with the figure
    table.set_figure(fig)
    plt.savefig('table.png', bbox_inches='tight', dpi=300)
    
    # Close the plot
    plt.close()

def save_barchart(data, handicap_tuple):

    data = data[data["Team"]!="Team"]
    teams = data["Team"].unique()
    team1 = teams[0]
    team2 = teams[1]
    data1 = data[data["Team"]==team1]
    data2 = data[data["Team"]==team2]

    goals1 = len(data1["Goal"].dropna())
    goals2 = len(data2["Goal"].dropna())

    fouls1 = len(data1[data1["Foul"]=="Foul"])
    fouls2 = len(data2[data2["Foul"]=="Foul"])

    shotAtGoal1 = len(data1["Shot At Goal"].dropna())
    shotAtGoal2 = len(data2["Shot At Goal"].dropna())

    shotAtGoalOutcomeGoal1 = len(data1[data1["Shot At Goal Outcome"]=="Goal"])
    shotAtGoalOutcomeGoal1 += len(data1[data1["Shot At Goal Outcome"]=="2pt Goal"])*2
    shotAtGoalOutcomeGoal2 = len(data2[data2["Shot At Goal Outcome"]=="Goal"])
    shotAtGoalOutcomeGoal2 += len(data2[data2["Shot At Goal Outcome"]=="2pt Goal"])*2

    penaltyShots1 = len(data1["Penalty Outcome"].dropna())
    penaltyShots2 = len(data2["Penalty Outcome"].dropna())

    penaltyShotsGoal1 = len(data1[data1["Penalty Outcome"]=="Goal"])
    penaltyShotsGoal2 = len(data2[data2["Penalty Outcome"]=="Goal"])

    throw_ins = data['Throw-in Outcome'].dropna()
    throw_ins_1 = throw_ins.str.contains(team1).sum()
    throw_ins_2 = throw_ins.str.contains(team2).sum()

    #### Handicap data ####
    handicap_teams = [team1,team2]
    handicap_goals = [0,0]
    data['Goal_New'] = data['Goal'].map(f)
    filter = data['Goal_New'].str.contains('handicap') == True
    handicap_goal_sum = sum(filter)
    if handicap_goal_sum>0:
        handicap_goals_team = data[filter]['Team'].iloc[0]
        if team1 == handicap_goals_team:
            handicap_goals[0] = handicap_goal_sum
        else:
            handicap_goals[1] = handicap_goal_sum    
    #### Handicap data end####
    
    ###############################
    ###############################
    hg_winning, hg_losing = handicap_tuple
    hg_winning = float(hg_winning)
    if hg_winning.is_integer():
        hg_winning = int(hg_winning)
    hg_losing = float(hg_losing)
    if hg_losing.is_integer():
        hg_losing = int(hg_losing)

    if goals2>goals1:
        handicap_goals[1] += hg_winning
        handicap_goals[0] += hg_losing
    else:
        handicap_goals[0] += hg_winning
        handicap_goals[1] += hg_losing
    
    # Convert handicap goals to integers if they are whole numbers
    if isinstance(handicap_goals[0], float) and handicap_goals[0].is_integer():
        handicap_goals[0] = int(handicap_goals[0])
    if isinstance(handicap_goals[1], float) and handicap_goals[1].is_integer():
        handicap_goals[1] = int(handicap_goals[1])

    handicap_goal_sum = handicap_goals[0] + handicap_goals[1]
    ###############################
    ###############################

    if handicap_goal_sum>0:
        # Data
        categories = ['Goals', 'Field Shots', 'Field Goals', 'Penalty Shots', 'Penalty Goals', 
                      'Throw-ins Won', 'Fouls Committed', 'Handicap Goals']
        col1_values = [goals1, shotAtGoal1, shotAtGoalOutcomeGoal1, penaltyShots1, penaltyShotsGoal1, 
                       throw_ins_1, fouls1, handicap_goals[0]]
        col2_values = [goals2, shotAtGoal2, shotAtGoalOutcomeGoal2, penaltyShots2, penaltyShotsGoal2, 
                       throw_ins_2, fouls2, handicap_goals[1]]
    else:
        # Data
        categories = ['Goals', 'Field Shots', 'Field Goals', 'Penalty Shots', 'Penalty Goals', 'Throw-ins Won', 'Fouls Committed']
        col1_values = [goals1, shotAtGoal1, shotAtGoalOutcomeGoal1, penaltyShots1, penaltyShotsGoal1, throw_ins_1, fouls1]
        col2_values = [goals2, shotAtGoal2, shotAtGoalOutcomeGoal2, penaltyShots2, penaltyShotsGoal2, throw_ins_2, fouls2]

    ###############################
    
    # Set up the figure and axis
    fig, ax = plt.subplots(figsize=(10, 4))  # Increase the figure width

    # Reverse the order of the data
    categories = categories[::-1]
    col1_values = col1_values[::-1]
    col2_values = col2_values[::-1]

    # Generate positions for the bars on the y-axis
    y = np.arange(len(categories))

    # Normalize the data
    total_values = np.add(col1_values, col2_values)
    normalized_col1 = np.divide(col1_values, total_values)
    normalized_col2 = np.divide(col2_values, total_values)

    # Define the bar width
    bar_width = 0.5  # Adjust the bar width as desired

    # Plot the bars
    bar1 = ax.barh(y, normalized_col1, height=bar_width, color='#800000', label='Scone Polo')  # Maroon red color
    bar2 = ax.barh(y, normalized_col2, height=bar_width, color='#808080', left=normalized_col1, label='Black Bears')  # Grey color

    # Remove x-axis ticks and labels
    ax.set_xticks([])

    # Remove ticks and lines
    ax.tick_params(axis='both', which='both', bottom=False, top=False, left=False, right=False)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    # Set the y-axis labels
    ax.set_yticks(y)
    ax.set_yticklabels(categories, fontsize=14)

    # Add numbers on the bars
    for i, (bar1, bar2) in enumerate(zip(bar1, bar2)):
        width1 = bar1.get_width()
        width2 = bar2.get_width()
        if col1_values[i]!=0:
            ax.text(width1 / 2, bar1.get_y() + bar1.get_height() / 2, str(col1_values[i]), ha='center', va='center', color='white', fontsize=14)
        if col2_values[i]!=0:
            ax.text(width1 + width2 / 2, bar2.get_y() + bar2.get_height() / 2, str(col2_values[i]), ha='center', va='center', color='white', fontsize=14)

    # Add legend
    legend_handles = [plt.Rectangle((0, 0), 1, 1, color='#800000'), plt.Rectangle((0, 0), 1, 1, color='#808080')]
    legend_labels = [team1, team2]
    ax.legend(legend_handles, legend_labels, loc='lower center', ncol=2, bbox_to_anchor=(0.5, -0.2), fontsize=14)

    # Save the figure as PNG
    plt.tight_layout()
    plt.savefig('stacked_bar_chart.png', dpi=300, bbox_inches='tight');  # Adjust the filename and DPI as desired

    # Close the plot
    plt.close()

def save_line(data, handicap_tuple):

    data = data[data["Team"]!="Team"]
    teams = data["Team"].unique()
    team1 = teams[0]
    team2 = teams[1]
    data1 = data[data["Team"]==team1]
    data2 = data[data["Team"]==team2]

    total_chukka = len(data['Chukka'].unique())
    
    #### Handicap data ####
    handicap_teams = [team1,team2]
    handicap_goals = [0,0]
    data['Goal_New'] = data['Goal'].map(f)
    filter = data['Goal_New'].str.contains('handicap') == True
    handicap_goal_sum = sum(filter)
    if handicap_goal_sum>0:
        handicap_goals_team = data[filter]['Team'].iloc[0]
        if team1 == handicap_goals_team:
            handicap_goals[0] = handicap_goal_sum
        else:
            handicap_goals[1] = handicap_goal_sum    
    #### Handicap data end####
    
    data1_goal = data1.loc[data1["Goal"].dropna().index]
    goals1Chukka1 = len(data1_goal[data1_goal["Chukka"]=="Chukka 1"])
    goals1Chukka2 = len(data1_goal[data1_goal["Chukka"]=="Chukka 2"])
    goals1Chukka3 = len(data1_goal[data1_goal["Chukka"]=="Chukka 3"])
    goals1Chukka4 = len(data1_goal[data1_goal["Chukka"]=="Chukka 4"])
    goals1Chukka5 = len(data1_goal[data1_goal["Chukka"]=="Chukka 5"])
    goals1Chukka6 = len(data1_goal[data1_goal["Chukka"]=="Chukka 6"])

    data2_goal = data2.loc[data2["Goal"].dropna().index]
    goals2Chukka1 = len(data2_goal[data2_goal["Chukka"]=="Chukka 1"])
    goals2Chukka2 = len(data2_goal[data2_goal["Chukka"]=="Chukka 2"])
    goals2Chukka3 = len(data2_goal[data2_goal["Chukka"]=="Chukka 3"])
    goals2Chukka4 = len(data2_goal[data2_goal["Chukka"]=="Chukka 4"])
    goals2Chukka5 = len(data2_goal[data2_goal["Chukka"]=="Chukka 5"])
    goals2Chukka6 = len(data2_goal[data2_goal["Chukka"]=="Chukka 6"])

    total_goals_1 = goals1Chukka1 + goals1Chukka2 + goals1Chukka3 + goals1Chukka4 + goals1Chukka5 + goals1Chukka6
    total_goals_2 = goals2Chukka1 + goals2Chukka2 + goals2Chukka3 + goals2Chukka4 + goals2Chukka5 + goals2Chukka6

    ###############################
    hg_winning, hg_losing = handicap_tuple
    hg_winning = float(hg_winning)
    if hg_winning.is_integer():
        hg_winning = int(hg_winning)
    hg_losing = float(hg_losing)
    if hg_losing.is_integer():
        hg_losing = int(hg_losing)

    if total_goals_2>total_goals_1:
        goals2Chukka1 += hg_winning
        goals1Chukka1 += hg_losing
    else:
        goals1Chukka1 += hg_winning
        goals2Chukka1 += hg_losing
    
    # Convert to integers if the results are whole numbers
    if isinstance(goals1Chukka1, float) and goals1Chukka1.is_integer():
        goals1Chukka1 = int(goals1Chukka1)
    if isinstance(goals2Chukka1, float) and goals2Chukka1.is_integer():
        goals2Chukka1 = int(goals2Chukka1)
    ###############################

    if (total_chukka == 4):
        # Create a DataFrame with column names and two rows
        data = pd.DataFrame(columns=['Team', 'Chukka1', 'Chukka2', 'Chukka3', 'Chukka4'])
        data.loc[0] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4]
        data.loc[1] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4]

        # Define Chukka names
        chukka_names = ['Chukka 1', 'Chukka 2', 'Chukka 3', 'Chukka 4']        

    elif (total_chukka == 7):
        goals1Chukka7 = len(data1_goal[data1_goal["Chukka"]=="Chukka 7"])
        goals2Chukka7 = len(data2_goal[data2_goal["Chukka"]=="Chukka 7"])
    
        # Create a DataFrame with column names and two rows
        data = pd.DataFrame(columns=['Team', 'Chukka1', 'Chukka2', 'Chukka3', 'Chukka4', 'Chukka5', 'Chukka6', 'Chukka7'])
        data.loc[0] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4, goals1Chukka5, goals1Chukka6, goals1Chukka7]
        data.loc[1] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4, goals2Chukka5, goals2Chukka6, goals2Chukka7]

        # Define Chukka names
        chukka_names = ['Chukka 1', 'Chukka 2', 'Chukka 3', 'Chukka 4', 'Chukka 5', 'Chukka 6', 'Chukka 7']

    elif (total_chukka == 8):
        goals1Chukka7 = len(data1_goal[data1_goal["Chukka"]=="Chukka 7"])
        goals2Chukka7 = len(data2_goal[data2_goal["Chukka"]=="Chukka 7"])
        goals1Chukka8 = len(data1_goal[data1_goal["Chukka"]=="Chukka 8"])
        goals2Chukka8 = len(data2_goal[data2_goal["Chukka"]=="Chukka 8"])
    
        # Create a DataFrame with column names and two rows
        data = pd.DataFrame(columns=['Team', 'Chukka1', 'Chukka2', 'Chukka3', 'Chukka4', 'Chukka5', 'Chukka6', 'Chukka7', 'Chukka8'])
        data.loc[0] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4, goals1Chukka5, goals1Chukka6, goals1Chukka7, goals1Chukka8]
        data.loc[1] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4, goals2Chukka5, goals2Chukka6, goals2Chukka7, goals2Chukka8]

        # Define Chukka names
        chukka_names = ['Chukka 1', 'Chukka 2', 'Chukka 3', 'Chukka 4', 'Chukka 5', 'Chukka 6', 'Chukka 7', 'Chukka 8']

    elif (total_chukka == 9):
        goals1Chukka7 = len(data1_goal[data1_goal["Chukka"]=="Chukka 7"])
        goals2Chukka7 = len(data2_goal[data2_goal["Chukka"]=="Chukka 7"])
        goals1Chukka8 = len(data1_goal[data1_goal["Chukka"]=="Chukka 8"])
        goals2Chukka8 = len(data2_goal[data2_goal["Chukka"]=="Chukka 8"])
        goals1Chukka9 = len(data1_goal[data1_goal["Chukka"]=="Chukka 9"])
        goals2Chukka9 = len(data2_goal[data2_goal["Chukka"]=="Chukka 9"])
    
        # Create a DataFrame with column names and two rows
        data = pd.DataFrame(columns=['Team', 'Ch 1', 'Ch 2', 'Ch 3', 'Ch 4', 'Ch 5', 'Ch 6', 'Ch 7', 'Ch 8', 'Ch 9'])
        data.loc[0] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4, goals1Chukka5, goals1Chukka6, goals1Chukka7, goals1Chukka8, goals1Chukka9]
        data.loc[1] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4, goals2Chukka5, goals2Chukka6, goals2Chukka7, goals2Chukka8, goals2Chukka9]

        # Define Chukka names
        chukka_names = ['Ch 1', 'Ch 2', 'Ch 3', 'Ch 4', 'Ch 5', 'Ch 6', 'Ch 7', 'Ch 8', 'Ch 9']

    else:        
        # Create a DataFrame with column names and two rows
        data = pd.DataFrame(columns=['Team', 'Chukka1', 'Chukka2', 'Chukka3', 'Chukka4', 'Chukka5', 'Chukka6'])
        data.loc[0] = [team1, goals1Chukka1, goals1Chukka2, goals1Chukka3, goals1Chukka4, goals1Chukka5, goals1Chukka6]
        data.loc[1] = [team2, goals2Chukka1, goals2Chukka2, goals2Chukka3, goals2Chukka4, goals2Chukka5, goals2Chukka6]

        # Define Chukka names
        chukka_names = ['Chukka 1', 'Chukka 2', 'Chukka 3', 'Chukka 4', 'Chukka 5', 'Chukka 6']


    # Cumulative sum
    data_cumsum = data.copy()
    data_cumsum.iloc[:, 1:] = data_cumsum.iloc[:, 1:].cumsum(axis=1)

    # Plotting
    plt.figure(figsize=(10, 6))
    
    # Define custom color codes
    maroon_red = '#800000'
    light_grey = '#D3D3D3'
    
    # Define colors for the lines
    colors = [maroon_red, light_grey]  # Add more colors if needed
    
    # Plot the lines and annotate points
    for index, row in data_cumsum.iterrows():
        x_values = row.index[1:]
        y_values = row.values[1:]

        plt.plot(x_values, y_values, marker='o', label=row['Team'], color=colors[index])

        # Annotate each point with its y-axis value for both lines
        for i, txt in enumerate(y_values):
            txt = float(txt)
            if txt.is_integer():
                txt = int(txt)

            # Skip annotation if the value is 0
            if txt == 0:
                continue

            # Determine the offset based on comparison with the other line
            other_index = 1 - index  # If index is 0, other_index is 1, and vice versa
            other_y_values = data_cumsum.iloc[other_index].values[1:]

            if txt > other_y_values[i]:
                offset = 10
            else:
                offset = -20

            plt.annotate(txt, (x_values[i], y_values[i]), textcoords="offset points", xytext=(2, offset), ha='center', fontsize=14)

    # Adjust the y-axis to add padding
    y_min, y_max = plt.ylim()
    padding = (y_max - y_min) * 0.05  # Add 5% padding on top and bottom
    plt.ylim(y_min - padding, y_max + padding)

    plt.xticks(range(0, len(chukka_names)), chukka_names, fontsize=16)
    plt.yticks(range(0, int(data_cumsum.iloc[:, 1:].values.max()) + 3, 2), fontsize=16)
    plt.legend()
    plt.grid(axis='y')
    plt.gca().spines['top'].set_visible(False)
    plt.gca().spines['right'].set_visible(False)
    plt.gca().spines['left'].set_visible(False)
    plt.gca().spines['bottom'].set_visible(False)
    plt.gca().tick_params(axis='x', which='both', bottom=False, top=False)
    plt.gca().tick_params(axis='y', which='both', left=False, right=False)
    plt.legend(loc='lower center', bbox_to_anchor=(0.5, -0.2), shadow=True, ncol=2, fontsize=16)
    plt.tight_layout()
    #plt.show()

    plt.savefig('stacked_line_chart.png', dpi=300, bbox_inches='tight');  # Adjust the filename and DPI as desired

    # Close the plot
    plt.close()
    
def textsize(text, font):
    im = Image.new(mode="P", size=(0, 0))
    draw = ImageDraw.Draw(im)
    _, _, width, height = draw.textbbox((0, 0), text=text, font=font)
    return width, height

def load_font_with_fallback(font_size):
    """
    Load a TTF font with graceful fallbacks across platforms.
    """
    try:
        return ImageFont.truetype("arial.ttf", font_size)
    except OSError:
        for path in [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        ]:
            try:
                return ImageFont.truetype(path, font_size)
            except OSError:
                continue
        return ImageFont.load_default()

def render_centered_text_image(
    text,
    base_width,
    height,
    background_color,
    text_color,
    initial_font_size,
    min_font_size=14,
    horizontal_padding=40,
):
    """
    Render a single-line centered text into an image. If the text does not fit the
    base_width, reduce font size down to min_font_size; if still too wide, expand
    the canvas width to fit with horizontal padding.
    """
    font_size = int(initial_font_size)
    font = load_font_with_fallback(font_size)
    text_width, text_height = textsize(text, font)
    width = int(base_width)

    # Try to shrink font to fit into current width
    while text_width + horizontal_padding > width and font_size > min_font_size:
        font_size -= 1
        font = load_font_with_fallback(font_size)
        text_width, text_height = textsize(text, font)

    # If still doesn't fit, expand width to fit the text
    if text_width + horizontal_padding > width:
        width = text_width + horizontal_padding

    image = Image.new("RGB", (width, height), background_color)
    draw = ImageDraw.Draw(image)
    position = ((width - text_width) // 2, (height - text_height) // 2)
    draw.text(position, text, fill=text_color, font=font)
    return image

def render_single_line_banner(
    text,
    target_width_px,
    background_color='#800000',
    text_color='#ffffff',
    max_font_size=120,
    min_font_size=28,
    horizontal_padding=40,
    vertical_padding=20,
):
    """
    Create a single-line banner image where the text fills the target width.
    The function chooses the largest font size that keeps the text (plus padding)
    within target_width_px, and uses minimal vertical padding for a tight banner.
    """
    font_size = int(max_font_size)
    font = load_font_with_fallback(font_size)
    text_width, text_height = textsize(text, font)

    while (text_width + 2 * horizontal_padding) > target_width_px and font_size > min_font_size:
        font_size -= 1
        font = load_font_with_fallback(font_size)
        text_width, text_height = textsize(text, font)

    width = max(target_width_px, text_width + 2 * horizontal_padding)
    height = text_height + 2 * vertical_padding

    image = Image.new("RGB", (width, height), background_color)
    draw = ImageDraw.Draw(image)
    position = ((width - text_width) // 2, (height - text_height) // 2)
    draw.text(position, text, fill=text_color, font=font)
    return image

def save_team_goals(data,data_file_filename, handicap_tuple):
    finalLst = data_Extract(data)
    extracted_data, filename = single_file(finalLst, data_file_filename)
    team1_name = extracted_data['team1'][0]
    team2_name = extracted_data['team2'][0]
    team1_goals = extracted_data['goals1'][0]
    team2_goals = extracted_data['goals2'][0]

    # Base styling for score line image
    base_width, height = 1000, 100
    bg_color = '#800000'  # Maroon Red Color

    hg_winning, hg_losing = handicap_tuple

    hg_winning = float(hg_winning)
    if hg_winning.is_integer():
        hg_winning = int(hg_winning)

    hg_losing = float(hg_losing)
    if hg_losing.is_integer():
        hg_losing = int(hg_losing)

    if team2_goals > team1_goals:
        team2_total = team2_goals + hg_winning
        team1_total = team1_goals + hg_losing
        # Convert to int if the result is a whole number
        if isinstance(team2_total, float) and team2_total.is_integer():
            team2_total = int(team2_total)
        if isinstance(team1_total, float) and team1_total.is_integer():
            team1_total = int(team1_total)
        text = f"{team2_name}          {team2_total}   -    {team1_total}            {team1_name}"
    else:
        team1_total = team1_goals + hg_winning
        team2_total = team2_goals + hg_losing
        # Convert to int if the result is a whole number
        if isinstance(team1_total, float) and team1_total.is_integer():
            team1_total = int(team1_total)
        if isinstance(team2_total, float) and team2_total.is_integer():
            team2_total = int(team2_total)
        text = f"{team1_name}          {team1_total}   -    {team2_total}            {team2_name}"

    # Render centered text, fitting or expanding width as needed
    image = render_centered_text_image(
        text=text,
        base_width=base_width,
        height=height,
        background_color=bg_color,
        text_color='#ffffff',
        initial_font_size=25,
        min_font_size=14,
        horizontal_padding=40,
    )

    # Save the image
    image.save("formatted_result.png")

def save_game_info(data, data_file_filename):
    finalLst = data_Extract(data)
    extracted_data, filename = single_file(finalLst, data_file_filename)
    
    # Parse date flexibly
    input_date_str = extracted_data['date'][0].strip()
    for fmt in ("%B %d. %Y", "%B %d %Y", "%B %d.%Y"):
        try:
            input_date = datetime.strptime(input_date_str, fmt)
            break
        except ValueError:
            continue
    else:
        raise ValueError(
            f"Date '{input_date_str}' is not in a supported format "
            "(expected like 'January 17. 2021' or 'January 17 2021')"
        )

    # Format the datetime object to the desired output format
    output_date_str = input_date.strftime("%Y-%m-%d")

    club = extracted_data['club'][0]
    tournament = extracted_data['tournament'][0]
    
    # Prepare text and render into an image that fits
    text = f"{output_date_str} - {club} - {tournament}"
    image = render_centered_text_image(
        text=text,
        base_width=1000,
        height=50,
        background_color='#800000',
        text_color='#ffffff',
        initial_font_size=25,
        min_font_size=14,
        horizontal_padding=40,
    )

    # Save the image
    image.save("formatted_result_game.png")

def get_text_first_paragraph(data, data_file_filename, toggle_button='other'):

    finalLst = data_Extract(data)
    extracted_data, filename = single_file(finalLst, data_file_filename)
    
    extracted_data_cols = list(extracted_data.columns)
    extracted_data_values = extracted_data[extracted_data["filename"] == filename].values[0]
    
    extracted_data_dict = {}
    for col, val in zip(extracted_data_cols, extracted_data_values):
        extracted_data_dict[col] = val
    
    winner = 'draw'
    if extracted_data_dict['goals1'] > extracted_data_dict['goals2']:
        winner = 'team1'
    elif extracted_data_dict['goals1'] < extracted_data_dict['goals2']:
        winner = 'team2'

    
    # Parse date flexibly
    input_date_str = extracted_data['date'][0].strip()
    for fmt in ("%B %d. %Y", "%B %d %Y", "%B %d.%Y"):
        try:
            input_date = datetime.strptime(input_date_str, fmt)
            break
        except ValueError:
            continue
    else:
        raise ValueError(
            f"Date '{input_date_str}' is not in a supported format "
            "(expected like 'January 17. 2021' or 'January 17 2021')"
        )

    weekday_string = input_date.strftime('%A')
    
    output = ''
    
    # --------------------
    if toggle_button == 'Semifinal':
        if winner == 'team1':
            output = f"{extracted_data_dict['team1']} won the Semi-final of the \
{extracted_data_dict['tournament']} on {weekday_string}, winning \
{extracted_data_dict['goals1']}-{extracted_data_dict['goals2']} against \
{extracted_data_dict['team2']} at the {extracted_data_dict['club']}"
    
        elif winner == 'team2':
            output = f"{extracted_data_dict['team2']} won the Semi-final of the \
{extracted_data_dict['tournament']} on {weekday_string}, winning \
{extracted_data_dict['goals2']}-{extracted_data_dict['goals1']} against \
{extracted_data_dict['team1']} at the {extracted_data_dict['club']}"
    
    # --------------------
    elif toggle_button == 'Final':
        if winner == 'team1':
            output = f"{extracted_data_dict['team1']} won the final of the \
{extracted_data_dict['tournament']} on {weekday_string}, winning \
{extracted_data_dict['goals1']}-{extracted_data_dict['goals2']} against \
{extracted_data_dict['team2']} at the {extracted_data_dict['club']}"
    
        elif winner == 'team2':
            output = f"{extracted_data_dict['team2']} won the final of the \
{extracted_data_dict['tournament']} on {weekday_string}, winning \
{extracted_data_dict['goals2']}-{extracted_data_dict['goals1']} against \
{extracted_data_dict['team1']} at the {extracted_data_dict['club']}"
    
    # --------------------
    elif toggle_button == 'Subsidiary final':
        if winner == 'team1':
            output = f"{extracted_data_dict['team1']} won the Subsidiary-final of the \
{extracted_data_dict['tournament']} on {weekday_string}, winning \
{extracted_data_dict['goals1']}-{extracted_data_dict['goals2']} against \
{extracted_data_dict['team2']} at the {extracted_data_dict['club']}"
    
        elif winner == 'team2':
            output = f"{extracted_data_dict['team2']} won the Subsidiary-final of the \
{extracted_data_dict['tournament']} on {weekday_string}, winning \
{extracted_data_dict['goals2']}-{extracted_data_dict['goals1']} against \
{extracted_data_dict['team1']} at the {extracted_data_dict['club']}"
    
    elif toggle_button == 'other':
        if winner == 'team1':
            output = f"{extracted_data_dict['team1']} won against \
{extracted_data_dict['team2']} by a score of \
{extracted_data_dict['goals1']}-{extracted_data_dict['goals2']} \
in the {extracted_data_dict['tournament']} at {extracted_data_dict['club']} \
on {weekday_string}"
    
        elif winner == 'team2':
            output = f"{extracted_data_dict['team2']} won against \
{extracted_data_dict['team1']} by a score of \
{extracted_data_dict['goals2']}-{extracted_data_dict['goals1']} \
in the {extracted_data_dict['tournament']} at {extracted_data_dict['club']} \
on {weekday_string}"

    return output

def get_header(data, data_file_filename):
    finalLst = data_Extract(data)
    extracted_data, filename = single_file(finalLst, data_file_filename)
    team1_name = extracted_data['team1'][0]
    team2_name = extracted_data['team2'][0]
    title = f"{team1_name} vs {team2_name}"
    
    # Parse date flexibly
    input_date_str = extracted_data['date'][0].strip()
    for fmt in ("%B %d. %Y", "%B %d %Y", "%B %d.%Y"):
        try:
            input_date = datetime.strptime(input_date_str, fmt)
            break
        except ValueError:
            continue
    else:
        raise ValueError(
            f"Date '{input_date_str}' is not in a supported format "
            "(expected like 'January 17. 2021' or 'January 17 2021')"
        )

    # Format the datetime object to the desired output format
    output_date_str = input_date.strftime("%Y-%m-%d")
    
    return title, output_date_str

def save_title_date(data, data_file_filename):
    title, output_date_str = get_header(data, data_file_filename)
    # Render a tight single-line banner with large font to minimize extra space
    banner = render_single_line_banner(
        text=title,
        target_width_px=2400,
        background_color='#FFFFFF',
        text_color='#000000',
        max_font_size=120,
        min_font_size=36,
        horizontal_padding=32,
        vertical_padding=12,
    )
    banner.save('match_info.png')

def create_combine_image():

    input_files = ['match_info.png', 'formatted_result_game.png', 'formatted_result.png', 'table.png', 'stacked_line_chart.png', 'stacked_bar_chart.png']
    
    # Open each input image
    images = [Image.open(file) for file in input_files]
    
    # Get the maximum width and total height of all images
    max_width = max(image.size[0] for image in images)
    total_height = sum(image.size[1] for image in images)
    
    # Specify the spacing between each image
    spacing = 50
    
    # Calculate the total height of the combined image with spacing
    combined_height = total_height + (len(images) - 1) * spacing
    
    # Create a new blank image with the maximum width and combined height
    combined_image = Image.new("RGBA", (max_width, combined_height), (255, 255, 255, 0))  # Transparent background
    
    # Paste each input image onto the combined image with spacing
    y_offset = 0
    for image in images:
        combined_image.paste(image, ((max_width - image.size[0]) // 2, y_offset))
        y_offset += image.size[1] + spacing
    
    # Save the combined image
    combined_image.save("combined_image.png")
    
def player_orders_with_arena(fields_tuple_1, fields_tuple_2, arena_toggled=False):
    """
    Generate a layout for player orders.
    If arena_toggled is True, display only 3 players and adjust the height dynamically.
    """
    if arena_toggled:
        # Arena mode: 3 players only
        field1, field2, field3, field5, field6, field7 = fields_tuple_1
        field11, field22, field33, field55, field66, field77 = fields_tuple_2

        # Define the fields (3 rows) - preserve original format
        fields_left = [f'1.   {field1}', f'2.   {field2}', f'3.   {field3}']
        fields_middle_left = [f'{field5}', f'{field6}', f'{field7}']
        # Calculate total while preserving original format
        try:
            total_left = sum([float(i) for i in fields_middle_left])
        except (ValueError, TypeError):
            total_left = 0

        fields_middle_right = [f'1.   {field11}', f'2.   {field22}', f'3.   {field33}']
        fields_right = [f'{field55}', f'{field66}', f'{field77}']
        # Calculate total while preserving original format
        try:
            total_right = sum([float(i) for i in fields_right])
        except (ValueError, TypeError):
            total_right = 0

        y_positions = [0.85, 0.6, 0.3]  # Adjusted for 3 rows
        total_y_position = 0.1  # Total positioned below the last row
        fig_height = 3  # Smaller figure height for 3 players
    else:
        # Default mode: 4 players
        field1, field2, field3, field4, field5, field6, field7, field8 = fields_tuple_1
        field11, field22, field33, field44, field55, field66, field77, field88 = fields_tuple_2

        fields_left = [f'1.   {field1}', f'2.   {field2}', f'3.   {field3}', f'4.   {field4}']
        fields_middle_left = [f'{field5}', f'{field6}', f'{field7}', f'{field8}']
        # Calculate total while preserving original format
        try:
            total_left = sum([float(i) for i in fields_middle_left])
        except (ValueError, TypeError):
            total_left = 0

        fields_middle_right = [f'1.   {field11}', f'2.   {field22}', f'3.   {field33}', f'4.   {field44}']
        fields_right = [f'{field55}', f'{field66}', f'{field77}', f'{field88}']
        # Calculate total while preserving original format
        try:
            total_right = sum([float(i) for i in fields_right])
        except (ValueError, TypeError):
            total_right = 0

        y_positions = [0.9, 0.7, 0.5, 0.3]  # 4 rows
        total_y_position = 0.1  # Total at the bottom
        fig_height = 3  # Standard figure height for 4 players

    # Convert totals to integers if they are whole numbers (e.g., 40.0 -> 40)
    if isinstance(total_left, float) and total_left.is_integer():
        total_left = int(total_left)
    if isinstance(total_right, float) and total_right.is_integer():
        total_right = int(total_right)

    # Create the figure with dynamic height
    plt.figure(figsize=(10, fig_height))

    # Define x-positions for the groups
    left_x = 0.01
    middle_left_x = 0.46
    middle_right_x = 0.54
    right_x = 0.99

    # Plot the text for the leftmost group
    for i, text in enumerate(fields_left):
        plt.text(left_x, y_positions[i], text, fontsize=16, ha='left', va='center')

    # Plot the text for the middle-left group
    for i, text in enumerate(fields_middle_left):
        plt.text(middle_left_x, y_positions[i], text, fontsize=16, ha='left', va='center')

    # Plot the text for the middle-right group
    for i, text in enumerate(fields_middle_right):
        plt.text(middle_right_x, y_positions[i], text, fontsize=16, ha='left', va='center')

    # Plot the text for the rightmost group
    for i, text in enumerate(fields_right):
        plt.text(right_x, y_positions[i], text, fontsize=16, ha='left', va='center')

    # Add totals at the adjusted position
    plt.text(middle_left_x, total_y_position, total_left, fontsize=16, ha='left', va='center')
    plt.text(right_x, total_y_position, total_right, fontsize=16, ha='left', va='center')

    # Hide axes and save image
    plt.axis('off')
    plt.savefig('fields_layout.png', bbox_inches='tight')
